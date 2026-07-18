"""
exportador_medidas_docx.py

Gera um documento .docx contendo SOMENTE as medidas DAX do modelo: para cada
medida, o nome seguido do código DAX completo — sem tabelas, colunas,
relacionamentos, Power Query ou qualquer outra seção da documentação.

Esta é uma saída independente da documentação em Markdown (gerador.py /
templates/*.md.j2): não gera um Markdown e converte, e não reaproveita o
Jinja2. Em vez disso, usa a mesma fonte de dados já extraída pelo restante da
ferramenta — `ModeloPowerBI.medidas` (ver extracao.py / modelos.py) — apenas
trocando a camada de saída para .docx. Assim, qualquer melhoria futura na
extração de medidas (extracao.py) já beneficia esta exportação automaticamente,
sem duplicar lógica.

Depois que o DOCX-base é montado com python-docx, o WordprocessingML é
processado em memória para aplicar syntax highlighting nos parágrafos de DAX.
"""

from __future__ import annotations

import copy
import io
import re
import zipfile
from dataclasses import dataclass
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Pt

from modelos import ModeloPowerBI

_FONTE_CODIGO = "Consolas"

_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "xml": "http://www.w3.org/XML/1998/namespace",
}

_W_NS = _NS["w"]
_XML_NS = _NS["xml"]


def _w_tag(local_name: str) -> str:
    return f"{{{_W_NS}}}{local_name}"


def _xml_attr(local_name: str) -> str:
    return f"{{{_XML_NS}}}{local_name}"


_TAG_P = _w_tag("p")
_TAG_R = _w_tag("r")
_TAG_RPR = _w_tag("rPr")
_TAG_T = _w_tag("t")
_TAG_BR = _w_tag("br")
_TAG_COLOR = _w_tag("color")
_TAG_RFONTS = _w_tag("rFonts")

_ATTR_VAL = _w_tag("val")
_ATTR_ASCII = _w_tag("ascii")
_ATTR_HANSI = _w_tag("hAnsi")
_ATTR_CS = _w_tag("cs")
_ATTR_EAST_ASIA = _w_tag("eastAsia")

_COLOR_PRETO = "000000"
_COLOR_FUNCAO = "0066CC"
_COLOR_TABELA = "1E3C96"
_COLOR_COLUNA = "0066CC"
_COLOR_MEDIDA = "9900CC"
_COLOR_STRING = "C41A16"
_COLOR_VARIAVEL = "00AAFF"
_COLOR_COMENTARIO = "008000"

_COLORIR_TEXTO_PADRAO = True

_DAX_FUNCTIONS = {
    "ABS",
    "ADDCOLUMNS",
    "ALL",
    "ALLSELECTED",
    "AVERAGE",
    "AVERAGEX",
    "BLANK",
    "CALCULATE",
    "CALCULATETABLE",
    "COALESCE",
    "CONCATENATE",
    "CONCATENATEX",
    "COUNT",
    "COUNTROWS",
    "CROSSFILTER",
    "DATE",
    "DATEADD",
    "DATESYTD",
    "DAY",
    "DISTINCT",
    "DISTINCTCOUNT",
    "DIVIDE",
    "ENDOFMONTH",
    "EOMONTH",
    "EXCEPT",
    "FALSE",
    "FILTER",
    "FORMAT",
    "GENERATE",
    "HASONEVALUE",
    "IF",
    "IN",
    "INTERSECT",
    "ISBLANK",
    "ISCROSSFILTERED",
    "ISFILTERED",
    "ISINSCOPE",
    "KEEPFILTERS",
    "LOOKUPVALUE",
    "MAX",
    "MAXX",
    "MIN",
    "MINX",
    "MONTH",
    "NOT",
    "NOW",
    "RANKX",
    "RELATED",
    "RELATEDTABLE",
    "REMOVEFILTERS",
    "RETURN",
    "SAMEPERIODLASTYEAR",
    "SELECTEDVALUE",
    "STARTOFMONTH",
    "SUM",
    "SUMMARIZE",
    "SUMMARIZECOLUMNS",
    "SUMX",
    "SWITCH",
    "TODAY",
    "TOPN",
    "TOTALMTD",
    "TOTALQTD",
    "TOTALYTD",
    "TRUE",
    "UNION",
    "VALUES",
    "VAR",
    "YEAR",
}

_IDENT = r"[^\W\d]\w*"
_FUNCTION_RE = re.compile(
    r"(?<!\w)("
    + "|".join(sorted(map(re.escape, _DAX_FUNCTIONS), key=len, reverse=True))
    + r")(?!\w)",
    re.IGNORECASE | re.UNICODE,
)
_VAR_DECL_RE = re.compile(
    r"(?<!\w)VAR(?!\w)(?P<gap>\s+)(?P<name>" + _IDENT + r")",
    re.IGNORECASE | re.UNICODE,
)
_QUOTED_TABLE_RE = re.compile(
    r"'(?P<table>[^'\r\n]+)'\s*(?P<bracket>\[(?P<column>[^\]\r\n]*)\])",
    re.UNICODE,
)
_UNQUOTED_TABLE_RE = re.compile(
    r"(?<![\w\]'\".])(?P<table>"
    + _IDENT
    + r")\s*(?P<bracket>\[(?P<column>[^\]\r\n]*)\])",
    re.UNICODE,
)
_MEASURE_RE = re.compile(r"\[(?P<measure>[^\]\r\n]*)\]", re.UNICODE)

ET.register_namespace("w", _W_NS)
ET.register_namespace("xml", _XML_NS)


@dataclass(frozen=True)
class _RunText:
    run: ET.Element
    text_node: ET.Element
    start: int
    end: int


@dataclass(frozen=True)
class _ParagraphText:
    paragraph: ET.Element
    text: str
    runs: list[_RunText]


def gerar_docx_medidas(modelo: ModeloPowerBI) -> bytes:
    """
    Gera, em memória, um documento .docx contendo apenas as medidas DAX de
    `modelo` e retorna seu conteúdo como bytes (pronto para uso em
    `st.download_button` ou para ser salvo em disco).

    Para cada medida: um parágrafo com o nome (negrito) seguido de um
    parágrafo com a expressão DAX completa (fonte monoespaçada e syntax
    highlighting). Se o modelo não tiver nenhuma medida, o documento gerado
    contém apenas uma mensagem indicando isso.
    """
    documento = Document()

    if not modelo.medidas:
        documento.add_paragraph("Nenhuma medida DAX encontrada neste modelo.")
    else:
        ultima = len(modelo.medidas) - 1
        for indice, medida in enumerate(modelo.medidas):
            _adicionar_medida(documento, medida.nome, medida.expressao)
            if indice != ultima:
                documento.add_paragraph()  # linha em branco separando as medidas

    buffer = io.BytesIO()
    documento.save(buffer)
    conteudo_docx = buffer.getvalue()
    if not modelo.medidas:
        return conteudo_docx

    return formatar_docx_medidas_dax(conteudo_docx)


def _adicionar_medida(documento: Document, nome: str, expressao_dax: str) -> None:
    paragrafo_nome = documento.add_paragraph()
    execucao_nome = paragrafo_nome.add_run(nome)
    execucao_nome.bold = True
    execucao_nome.font.size = Pt(12)

    # O código DAX pode ter múltiplas linhas; cada quebra de linha vira um
    # novo "run" quebrado (add_break), já que o docx não interpreta "\n".
    paragrafo_dax = documento.add_paragraph()
    linhas = (expressao_dax or "").split("\n")
    for indice_linha, linha in enumerate(linhas):
        execucao_dax = paragrafo_dax.add_run(linha)
        execucao_dax.font.name = _FONTE_CODIGO
        execucao_dax.font.size = Pt(10)
        if indice_linha != len(linhas) - 1:
            execucao_dax.add_break()


def formatar_docx_medidas_dax(conteudo_docx: bytes) -> bytes:
    """
    Aplica syntax highlighting DAX em um DOCX já gerado e retorna novos bytes.

    A formatação é aplicada apenas nos parágrafos que usam a fonte de código
    do exportador, preservando nomes das medidas, parágrafos em branco e todo
    o restante da estrutura do documento.
    """
    arquivos, arvore = _ler_documento_docx(conteudo_docx)
    texto_antes = _extrair_texto_documento(arvore)

    _processar_arvore(arvore)

    texto_depois = _extrair_texto_documento(arvore)
    if texto_antes != texto_depois:
        raise RuntimeError(
            "Falha de segurança: o texto concatenado mudou. O DOCX formatado não foi gerado."
        )

    return _salvar_documento_bytes(arquivos, arvore)


def _ler_documento_docx(conteudo_docx: bytes) -> tuple[dict[str, bytes], ET.ElementTree]:
    with zipfile.ZipFile(io.BytesIO(conteudo_docx), "r") as docx:
        arquivos = {info.filename: docx.read(info.filename) for info in docx.infolist()}

    if "word/document.xml" not in arquivos:
        raise ValueError("O arquivo .docx não contém word/document.xml.")

    parser = ET.XMLParser(target=ET.TreeBuilder(insert_comments=True, insert_pis=True))
    raiz = ET.fromstring(arquivos["word/document.xml"], parser=parser)
    return arquivos, ET.ElementTree(raiz)


def _iter_text_nodes_e_quebras(
    elemento: ET.Element,
    run_atual: ET.Element | None = None,
):
    if elemento.tag == _TAG_R:
        run_atual = elemento

    if elemento.tag == _TAG_T and run_atual is not None:
        yield run_atual, elemento
    elif elemento.tag == _TAG_BR and run_atual is not None:
        yield run_atual, elemento

    for child in list(elemento):
        yield from _iter_text_nodes_e_quebras(child, run_atual)


def _extrair_runs(paragrafo: ET.Element) -> list[_RunText]:
    runs: list[_RunText] = []
    posicao = 0

    for run, node in _iter_text_nodes_e_quebras(paragrafo):
        if node.tag == _TAG_BR:
            posicao += 1
            continue

        texto = node.text or ""
        inicio = posicao
        fim = inicio + len(texto)
        runs.append(_RunText(run=run, text_node=node, start=inicio, end=fim))
        posicao = fim

    return runs


def _reconstruir_texto(paragrafo: ET.Element) -> _ParagraphText:
    runs: list[_RunText] = []
    partes: list[str] = []
    posicao = 0

    for run, node in _iter_text_nodes_e_quebras(paragrafo):
        if node.tag == _TAG_BR:
            partes.append("\n")
            posicao += 1
            continue

        texto = node.text or ""
        inicio = posicao
        fim = inicio + len(texto)
        runs.append(_RunText(run=run, text_node=node, start=inicio, end=fim))
        partes.append(texto)
        posicao = fim

    return _ParagraphText(paragraph=paragrafo, text="".join(partes), runs=runs)


def _extrair_texto_documento(arvore: ET.ElementTree) -> str:
    return "\n".join(
        _reconstruir_texto(paragrafo).text for paragrafo in arvore.getroot().iter(_TAG_P)
    )


def _intervalo_livre(ocupado: list[bool], inicio: int, fim: int) -> bool:
    if inicio < 0 or fim > len(ocupado) or inicio >= fim:
        return False
    return not any(ocupado[inicio:fim])


def _marcar(ocupado: list[bool], inicio: int, fim: int) -> None:
    for i in range(inicio, fim):
        ocupado[i] = True


def _aplicar_cor_intervalo(
    cores: list[str | None],
    inicio: int,
    fim: int,
    cor: str,
) -> None:
    for i in range(inicio, fim):
        cores[i] = cor


def _scan_strings_e_comentarios(texto: str) -> tuple[list[tuple[int, int]], list[tuple[int, int]]]:
    strings: list[tuple[int, int]] = []
    comentarios: list[tuple[int, int]] = []
    i = 0
    n = len(texto)

    while i < n:
        char = texto[i]

        if char == '"':
            inicio = i
            i += 1
            while i < n:
                if texto[i] == '"':
                    if i + 1 < n and texto[i + 1] == '"':
                        i += 2
                        continue
                    i += 1
                    break
                i += 1
            strings.append((inicio, i))
            continue

        if char == "/" and i + 1 < n and texto[i + 1] == "/":
            inicio = i
            i += 2
            while i < n and texto[i] not in "\r\n":
                i += 1
            comentarios.append((inicio, i))
            continue

        i += 1

    return strings, comentarios


def _coletar_variaveis(textos: Iterable[str]) -> set[str]:
    variaveis: set[str] = set()

    for texto in textos:
        protegido = [False] * len(texto)
        strings, comentarios = _scan_strings_e_comentarios(texto)
        for inicio, fim in strings + comentarios:
            _marcar(protegido, inicio, fim)

        for match in _VAR_DECL_RE.finditer(texto):
            inicio, fim = match.span("name")
            if _intervalo_livre(protegido, inicio, fim):
                variaveis.add(match.group("name").upper())

    return variaveis


def _tokenizar_dax(texto: str, variaveis: set[str]) -> list[str | None]:
    cores: list[str | None] = [_COLOR_PRETO if _COLORIR_TEXTO_PADRAO else None for _ in texto]
    ocupado = [False] * len(texto)

    strings, comentarios = _scan_strings_e_comentarios(texto)

    for inicio, fim in strings:
        _aplicar_cor_intervalo(cores, inicio, fim, _COLOR_STRING)
        _marcar(ocupado, inicio, fim)

    for inicio, fim in comentarios:
        _aplicar_cor_intervalo(cores, inicio, fim, _COLOR_COMENTARIO)
        _marcar(ocupado, inicio, fim)

    _aplicar_tabelas_e_colunas(texto, cores, ocupado)
    _aplicar_medidas(texto, cores, ocupado)
    _aplicar_funcoes(texto, cores, ocupado)
    _aplicar_variaveis(texto, cores, ocupado, variaveis)

    return cores


def _aplicar_tabelas_e_colunas(
    texto: str,
    cores: list[str | None],
    ocupado: list[bool],
) -> None:
    for regex in (_QUOTED_TABLE_RE, _UNQUOTED_TABLE_RE):
        for match in regex.finditer(texto):
            inicio_ref, fim_ref = match.span()
            if not _intervalo_livre(ocupado, inicio_ref, fim_ref):
                continue

            tabela_inicio, tabela_fim = match.span("table")
            coluna_inicio, coluna_fim = match.span("column")
            _aplicar_cor_intervalo(cores, tabela_inicio, tabela_fim, _COLOR_TABELA)
            if coluna_inicio < coluna_fim:
                _aplicar_cor_intervalo(cores, coluna_inicio, coluna_fim, _COLOR_COLUNA)
            _marcar(ocupado, inicio_ref, fim_ref)


def _aplicar_medidas(texto: str, cores: list[str | None], ocupado: list[bool]) -> None:
    for match in _MEASURE_RE.finditer(texto):
        inicio_ref, fim_ref = match.span()
        if not _intervalo_livre(ocupado, inicio_ref, fim_ref):
            continue

        medida_inicio, medida_fim = match.span("measure")
        if medida_inicio < medida_fim:
            _aplicar_cor_intervalo(cores, medida_inicio, medida_fim, _COLOR_MEDIDA)
        _marcar(ocupado, inicio_ref, fim_ref)


def _aplicar_funcoes(texto: str, cores: list[str | None], ocupado: list[bool]) -> None:
    for match in _FUNCTION_RE.finditer(texto):
        inicio, fim = match.span()
        if not _intervalo_livre(ocupado, inicio, fim):
            continue

        _aplicar_cor_intervalo(cores, inicio, fim, _COLOR_FUNCAO)
        _marcar(ocupado, inicio, fim)


def _aplicar_variaveis(
    texto: str,
    cores: list[str | None],
    ocupado: list[bool],
    variaveis: set[str],
) -> None:
    if not variaveis:
        return

    for match in re.finditer(_IDENT, texto, re.UNICODE):
        nome = match.group(0).upper()
        inicio, fim = match.span()
        if nome not in variaveis or not _intervalo_livre(ocupado, inicio, fim):
            continue

        _aplicar_cor_intervalo(cores, inicio, fim, _COLOR_VARIAVEL)
        _marcar(ocupado, inicio, fim)


def _mapear_tokens_para_runs(
    paragraph_text: _ParagraphText,
    cores: list[str | None],
) -> dict[int, list[str | None]]:
    mapa: dict[int, list[str | None]] = {}
    for rt in paragraph_text.runs:
        mapa[id(rt.text_node)] = cores[rt.start : rt.end]
    return mapa


def _definir_cor_run(run: ET.Element, cor: str | None) -> None:
    if cor is None:
        return

    rpr = run.find(_TAG_RPR)
    if rpr is None:
        rpr = ET.Element(_TAG_RPR)
        run.insert(0, rpr)

    color = rpr.find(_TAG_COLOR)
    if color is None:
        color = ET.Element(_TAG_COLOR)
        rpr.append(color)

    color.set(_ATTR_VAL, cor)


def _novo_run_com_rpr(run_original: ET.Element, cor: str | None) -> ET.Element:
    novo = ET.Element(run_original.tag)
    for attr, valor in run_original.attrib.items():
        novo.set(attr, valor)

    rpr = run_original.find(_TAG_RPR)
    if rpr is not None:
        novo.append(copy.deepcopy(rpr))

    _definir_cor_run(novo, cor)
    return novo


def _ajustar_xml_space(text_node: ET.Element) -> None:
    texto = text_node.text or ""
    if texto != texto.strip() or "  " in texto or "\t" in texto:
        text_node.set(_xml_attr("space"), "preserve")


def _dividir_texto_por_cor(texto: str, cores: list[str | None]) -> list[tuple[str, str | None]]:
    if not texto:
        return []

    segmentos: list[tuple[str, str | None]] = []
    inicio = 0
    cor_atual = cores[0] if cores else None

    for i in range(1, len(texto)):
        if cores[i] != cor_atual:
            segmentos.append((texto[inicio:i], cor_atual))
            inicio = i
            cor_atual = cores[i]

    segmentos.append((texto[inicio:], cor_atual))
    return segmentos


def _cores_do_run(run: ET.Element, mapa: dict[int, list[str | None]]) -> list[str | None]:
    cores: list[str | None] = []
    for text_node in run.iter(_TAG_T):
        cores.extend(mapa.get(id(text_node), []))
    return cores


def _montar_parent_map(raiz: ET.Element) -> dict[int, ET.Element]:
    return {id(child): parent for parent in raiz.iter() for child in list(parent)}


def _aplicar_cores(paragrafo: ET.Element, mapa: dict[int, list[str | None]]) -> None:
    runs = list({id(rt.run): rt.run for rt in _extrair_runs(paragrafo)}.values())
    parent_map = _montar_parent_map(paragrafo)

    for run in runs:
        text_nodes = list(run.iter(_TAG_T))
        if not text_nodes:
            continue

        todas_cores = _cores_do_run(run, mapa)
        cores_unicas = {cor for cor in todas_cores}

        if len(cores_unicas) <= 1:
            _definir_cor_run(run, next(iter(cores_unicas), None))
            continue

        _substituir_run_por_segmentos(run, mapa, parent_map)


def _substituir_run_por_segmentos(
    run: ET.Element,
    mapa: dict[int, list[str | None]],
    parent_map: dict[int, ET.Element],
) -> None:
    pai = parent_map.get(id(run))
    if pai is None:
        return

    indice = list(pai).index(run)
    novos_runs: list[ET.Element] = []
    run_atual: ET.Element | None = None
    cor_atual: str | None = None

    def iniciar_ou_reusar_run(cor: str | None) -> ET.Element:
        nonlocal run_atual, cor_atual
        if run_atual is None or cor != cor_atual:
            run_atual = _novo_run_com_rpr(run, cor)
            novos_runs.append(run_atual)
            cor_atual = cor
        return run_atual

    for child in run:
        if child.tag == _TAG_RPR:
            continue

        if child.tag != _TAG_T:
            destino = iniciar_ou_reusar_run(cor_atual)
            destino.append(copy.deepcopy(child))
            continue

        texto = child.text or ""
        cores = mapa.get(id(child), [None] * len(texto))
        for segmento, cor in _dividir_texto_por_cor(texto, cores):
            if not segmento:
                continue

            destino = iniciar_ou_reusar_run(cor)
            novo_texto = copy.deepcopy(child)
            novo_texto.text = segmento
            _ajustar_xml_space(novo_texto)
            destino.append(novo_texto)

    if not novos_runs:
        return

    for novo in novos_runs:
        pai.insert(indice, novo)
        indice += 1

    pai.remove(run)


def _processar_arvore(arvore: ET.ElementTree) -> None:
    paragrafos = [_reconstruir_texto(p) for p in arvore.getroot().iter(_TAG_P)]
    paragrafos_dax = [pt for pt in paragrafos if _eh_paragrafo_dax(pt)]
    variaveis = _coletar_variaveis(pt.text for pt in paragrafos_dax)

    for pt in paragrafos_dax:
        if not pt.text:
            continue
        cores = _tokenizar_dax(pt.text, variaveis)
        mapa = _mapear_tokens_para_runs(pt, cores)
        _aplicar_cores(pt.paragraph, mapa)


def _eh_paragrafo_dax(paragraph_text: _ParagraphText) -> bool:
    if not paragraph_text.text:
        return False
    return any(_run_usa_fonte_codigo(rt.run) for rt in paragraph_text.runs)


def _run_usa_fonte_codigo(run: ET.Element) -> bool:
    rpr = run.find(_TAG_RPR)
    if rpr is None:
        return False

    rfonts = rpr.find(_TAG_RFONTS)
    if rfonts is None:
        return False

    fontes = {
        rfonts.get(_ATTR_ASCII),
        rfonts.get(_ATTR_HANSI),
        rfonts.get(_ATTR_CS),
        rfonts.get(_ATTR_EAST_ASIA),
    }
    return _FONTE_CODIGO in fontes


def _salvar_documento_bytes(arquivos: dict[str, bytes], arvore: ET.ElementTree) -> bytes:
    arquivos = dict(arquivos)
    buffer_xml = io.BytesIO()
    arvore.write(buffer_xml, encoding="UTF-8", xml_declaration=True, short_empty_elements=True)
    arquivos["word/document.xml"] = buffer_xml.getvalue()

    buffer_docx = io.BytesIO()
    with zipfile.ZipFile(buffer_docx, "w", compression=zipfile.ZIP_DEFLATED) as docx:
        for nome, conteudo in arquivos.items():
            docx.writestr(nome, conteudo)

    return buffer_docx.getvalue()
